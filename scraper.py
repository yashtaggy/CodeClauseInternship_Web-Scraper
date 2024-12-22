import requests
from bs4 import BeautifulSoup
from docx import Document

def scrape_website(url):
    try:
        # Fetch website content
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Create a new Document
            doc = Document()
            
            # Extract Headings (h1, h2, h3)
            headings = []
            for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                headings.append(heading.get_text(strip=True))
            
            # Extract Paragraphs
            paragraphs = []
            for para in soup.find_all('p'):
                paragraphs.append(para.get_text(strip=True))
            
            # Extract Links
            links = []
            for link in soup.find_all('a', href=True):
                links.append(link['href'])

            # Add content to the DOCX file
            doc.add_heading('Headings', level=1)
            if headings:
                for heading in headings:
                    doc.add_paragraph(heading)
            else:
                doc.add_paragraph("No headings found.")
            
            doc.add_heading('Paragraphs', level=1)
            if paragraphs:
                for para in paragraphs[:10]:  # Limit to first 10 paragraphs
                    doc.add_paragraph(para)
            else:
                doc.add_paragraph("No paragraphs found.")
            
            doc.add_heading('Links', level=1)
            if links:
                for link in links[:10]:  # Limit to first 10 links
                    doc.add_paragraph(link)
            else:
                doc.add_paragraph("No links found.")
            
            # Save the document
            doc.save('scraped_output.docx')
            print("Scraped data saved to 'scraped_output.docx'")
        else:
            print(f"Failed to fetch the website. Status code: {response.status_code}")
    except Exception as e:
        print(f"An error occurred: {e}")

# Input: URL
url = input("Enter the website URL: ").strip()
scrape_website(url)
